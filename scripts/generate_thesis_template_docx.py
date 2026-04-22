from __future__ import annotations

import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.table import _Cell
from docx.text.paragraph import Paragraph


XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
FIGURE_INSERTS = {
    "3.3.1 总体架构": [
        {
            "file": "fig_3_1_architecture.png",
            "intro": "系统总体架构如图3-1所示。",
            "caption": "图3-1 系统总体架构图",
            "width_cm": 15.8,
        }
    ],
    "3.4 功能模块设计": [
        {
            "file": "fig_3_2_modules.png",
            "intro": "系统核心功能模块划分如图3-2所示。",
            "caption": "图3-2 系统功能模块图",
            "width_cm": 15.2,
        }
    ],
    "3.5.2 逻辑结构设计": [
        {
            "file": "fig_3_3_er.png",
            "intro": "围绕订单、菜品、聊天和健康管理抽取的核心实体关系如图3-3所示。",
            "caption": "图3-3 核心数据实体关系图",
            "width_cm": 15.5,
        }
    ],
    "4.1.2 推荐引擎实现": [
        {
            "file": "fig_4_1_recommendation.png",
            "intro": "推荐引擎的主要处理流程如图4-1所示。",
            "caption": "图4-1 个性化推荐引擎处理流程图",
            "width_cm": 15.5,
        }
    ],
    "4.1.3 订单管理模块": [
        {
            "file": "fig_4_2_order_flow.png",
            "intro": "订单处理与即时通知流程如图4-2所示。",
            "caption": "图4-2 订单处理与即时通知流程图",
            "width_cm": 15.5,
        }
    ],
    "4.2.2 用户端页面实现": [
        {
            "file": "fig_4_3_user_merchant_list.png",
            "intro": "用户端商家查找页面如图4-3所示。",
            "caption": "图4-3 用户端商家查找页面",
            "width_cm": 15.8,
        },
        {
            "file": "fig_4_4_user_calorie.png",
            "intro": "用户端卡路里统计页面如图4-4所示。",
            "caption": "图4-4 用户端卡路里统计页面",
            "width_cm": 15.8,
        },
    ],
    "4.2.3 商家端页面实现": [
        {
            "file": "fig_4_5_merchant_orders.png",
            "intro": "商家端订单管理页面如图4-5所示。",
            "caption": "图4-5 商家端订单管理页面",
            "width_cm": 15.8,
        },
        {
            "file": "fig_4_6_merchant_statistics.png",
            "intro": "商家端经营统计页面如图4-6所示。",
            "caption": "图4-6 商家端经营统计页面",
            "width_cm": 15.8,
        },
    ],
    "4.4.2 智能助手实现": [
        {
            "file": "fig_4_7_ai_assistant.png",
            "intro": "AI饮食助手页面如图4-7所示。",
            "caption": "图4-7 AI饮食助手页面",
            "width_cm": 15.8,
        }
    ],
    "5.4.2 性能测试": [
        {
            "file": "fig_5_1_performance.png",
            "intro": "关键性能指标优化效果如图5-1所示。",
            "caption": "图5-1 关键性能指标优化效果图",
            "width_cm": 15.0,
        }
    ],
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="按学校模板生成论文终稿")
    parser.add_argument("--template", required=True, help="模板 docx 路径")
    parser.add_argument("--source", required=True, help="论文 markdown 路径")
    parser.add_argument("--output", required=True, help="输出 docx 路径")
    parser.add_argument("--figure-dir", help="论文插图目录")
    return parser.parse_args()


def clean_text(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = re.sub(r"\s+$", "", text)
    return text


def read_markdown(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8").splitlines()


def extract_front_matter(lines: list[str]) -> dict[str, object]:
    title = ""
    meta: dict[str, str] = {}
    cn_abs: list[str] = []
    en_abs: list[str] = []
    cn_keywords = ""
    en_keywords = ""

    current: str | None = None
    paragraph_buffer: list[str] = []

    def flush_buffer() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer or current is None:
            paragraph_buffer = []
            return
        paragraph = clean_text(" ".join(s.strip() for s in paragraph_buffer if s.strip()))
        if paragraph:
            if current == "cn_abs":
                cn_abs.append(paragraph)
            elif current == "en_abs":
                en_abs.append(paragraph)
        paragraph_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            if not title:
                title = clean_text(stripped[2:].strip())
            if stripped == "# 第一章 绪论":
                flush_buffer()
                break
            continue

        if stripped.startswith("题目："):
            meta["题目"] = clean_text(stripped.split("：", 1)[1].strip())
            continue
        if stripped.startswith("院别："):
            meta["院别"] = clean_text(stripped.split("：", 1)[1].strip())
            continue
        if stripped.startswith("专业："):
            meta["专业"] = clean_text(stripped.split("：", 1)[1].strip())
            continue
        if stripped.startswith("学生姓名："):
            meta["姓名"] = clean_text(stripped.split("：", 1)[1].strip())
            continue
        if stripped.startswith("学号："):
            meta["学号"] = clean_text(stripped.split("：", 1)[1].strip())
            continue

        if stripped == "## 摘要":
            flush_buffer()
            current = "cn_abs"
            continue
        if stripped == "## Abstract":
            flush_buffer()
            current = "en_abs"
            continue
        if stripped.startswith("关键词："):
            flush_buffer()
            cn_keywords = clean_text(stripped)
            current = None
            continue
        if stripped.startswith("Keywords:"):
            flush_buffer()
            en_keywords = clean_text(stripped)
            current = None
            continue
        if stripped == "## 目录":
            flush_buffer()
            current = None
            continue

        if current in {"cn_abs", "en_abs"}:
            if stripped:
                paragraph_buffer.append(stripped)
            else:
                flush_buffer()

    flush_buffer()

    return {
        "title": meta.get("题目", title),
        "college": meta.get("院别", ""),
        "major": meta.get("专业", ""),
        "name": meta.get("姓名", ""),
        "student_id": meta.get("学号", ""),
        "cn_abs": cn_abs,
        "cn_keywords": cn_keywords,
        "en_abs": en_abs,
        "en_keywords": en_keywords,
    }


def parse_body_blocks(lines: list[str]) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    in_body = False
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    current_h1 = ""

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = clean_text(" ".join(s.strip() for s in paragraph_buffer if s.strip()))
        if text:
            blocks.append({"type": "paragraph", "text": text})
        paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows: list[list[str]] = []
        for raw in table_buffer:
            stripped = raw.strip()
            if not stripped.startswith("|"):
                continue
            cells = [clean_text(cell.strip()) for cell in stripped.strip("|").split("|")]
            if all(re.fullmatch(r"-{3,}", cell.replace(" ", "")) for cell in cells):
                continue
            rows.append(cells)
        if rows:
            blocks.append({"type": "table", "rows": rows})
        table_buffer = []

    for line in lines:
        stripped = line.strip()
        if not in_body:
            if stripped == "# 第一章 绪论":
                in_body = True
            else:
                continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_buffer.append(stripped)
            continue
        flush_table()

        if stripped.startswith("# "):
            flush_paragraph()
            current_h1 = clean_text(stripped[2:].strip())
            blocks.append({"type": "heading1", "text": current_h1})
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            blocks.append({"type": "heading2", "text": clean_text(stripped[3:].strip())})
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            blocks.append({"type": "heading3", "text": clean_text(stripped[4:].strip())})
            continue

        if current_h1 == "第七章 参考文献" and re.match(r"^\d+\.\s", stripped):
            flush_paragraph()
            blocks.append({"type": "paragraph", "text": clean_text(stripped)})
            continue

        if stripped:
            paragraph_buffer.append(stripped)
        else:
            flush_paragraph()

    flush_table()
    flush_paragraph()
    return blocks


def paragraph_has_section(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.sectPr is not None)


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def build_run(paragraph: Paragraph, text: str, template_rpr) -> None:
    run = OxmlElement("w:r")
    if template_rpr is not None:
        run.append(deepcopy(template_rpr))
    if text:
        text_el = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            text_el.set(XML_SPACE, "preserve")
        text_el.text = text
        run.append(text_el)
    paragraph._p.append(run)


def set_paragraph_text_like_template(paragraph: Paragraph, text: str) -> None:
    template_rpr = None
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            rpr = child.find(qn("w:rPr"))
            if rpr is not None:
                template_rpr = deepcopy(rpr)
                break
    clear_paragraph(paragraph)
    if not text:
        return
    lines = text.split("\n")
    for index, chunk in enumerate(lines):
        build_run(paragraph, chunk, template_rpr)
        if index != len(lines) - 1:
            br_run = OxmlElement("w:r")
            if template_rpr is not None:
                br_run.append(deepcopy(template_rpr))
            br = OxmlElement("w:br")
            br_run.append(br)
            paragraph._p.append(br_run)


def insert_paragraph_after(anchor: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    return paragraph


def insert_paragraph_after_element(element, parent, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, parent)
    if style:
        paragraph.style = style
    return paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_siblings_between(start_paragraph: Paragraph, end_paragraph: Paragraph) -> None:
    current = start_paragraph._p.getnext()
    while current is not None and current is not end_paragraph._p:
        next_node = current.getnext()
        current.getparent().remove(current)
        current = next_node


def get_paragraph(document: Document, index: int) -> Paragraph:
    return document.paragraphs[index]


def find_first_heading_index(document: Document) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 1":
            return index
    raise ValueError("模板中未找到正文起始标题")


def create_field_run(paragraph: Paragraph, instr: str, placeholder: str) -> None:
    begin = OxmlElement("w:r")
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin.append(begin_char)
    paragraph._p.append(begin)

    instr_run = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(XML_SPACE, "preserve")
    instr_text.text = instr
    instr_run.append(instr_text)
    paragraph._p.append(instr_run)

    separate = OxmlElement("w:r")
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate.append(separate_char)
    paragraph._p.append(separate)

    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = placeholder
    placeholder_run.append(placeholder_text)
    paragraph._p.append(placeholder_run)

    end = OxmlElement("w:r")
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end.append(end_char)
    paragraph._p.append(end)


def add_toc(anchor: Paragraph, end_paragraph: Paragraph) -> Paragraph:
    remove_siblings_between(anchor, end_paragraph)
    set_paragraph_text_like_template(anchor, "目 录")
    clear_paragraph(end_paragraph)
    toc_para = insert_paragraph_after(anchor)
    toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    create_field_run(toc_para, 'TOC \\o "1-3" \\h \\z \\u', "目录将在打开文档后更新")
    return toc_para


def set_cover(front: dict[str, object], document: Document) -> None:
    field_texts = {
        6: f"院    别：    {front['college']}     ",
        7: f"专    业：    {front['major']}   ",
        8: f"姓    名：        {front['name']}         ",
        9: f"学    号：     {front['student_id']}     ",
        10: "指导教师：        张倩         ",
        11: "日    期：     2026年5月        ",
    }
    for index, text in field_texts.items():
        set_paragraph_text_like_template(get_paragraph(document, index), text)

    set_paragraph_text_like_template(
        get_paragraph(document, 21),
        "                       论文作者签名：          日期：2026年 5月    日   ",
    )
    set_paragraph_text_like_template(
        get_paragraph(document, 29),
        "                    论文作者签名：            导师签名：                                         日期：2026年 5月    日   ",
    )

    if document.tables:
        cover_table = document.tables[0]
        set_cell_text_like_template(
            cover_table.cell(0, 1),
            split_cn_title(str(front["title"])),
        )
        set_cell_text_like_template(
            cover_table.cell(1, 1),
            "Design and Implementation of an Intelligent Catering System\nBased on SpringBoot and Multi-terminal Integration",
        )


def split_cn_title(title: str) -> str:
    if "智能餐饮系统" in title and "设计与实现" in title:
        return title.replace("智能餐饮系统设计与实现", "智能餐饮系统\n设计与实现")
    return title


def build_abstract(front: dict[str, object], document: Document) -> None:
    cn_title = split_cn_title(str(front["title"]))
    set_paragraph_text_like_template(get_paragraph(document, 31), cn_title)

    cn_paragraph_indexes = [33, 34, 35]
    cn_abs = list(front["cn_abs"])
    for index, paragraph_index in enumerate(cn_paragraph_indexes):
        text = cn_abs[index] if index < len(cn_abs) else ""
        set_paragraph_text_like_template(get_paragraph(document, paragraph_index), text)
        format_body_paragraph(get_paragraph(document, paragraph_index))
    clear_paragraph(get_paragraph(document, 36))
    set_paragraph_text_like_template(get_paragraph(document, 37), str(front["cn_keywords"]))
    format_body_paragraph(get_paragraph(document, 37))

    set_paragraph_text_like_template(get_paragraph(document, 38), str(front["title"]))
    en_paragraph_indexes = [40, 41, 42]
    en_abs = list(front["en_abs"])
    for index, paragraph_index in enumerate(en_paragraph_indexes):
        text = en_abs[index] if index < len(en_abs) else ""
        set_paragraph_text_like_template(get_paragraph(document, paragraph_index), text)
        format_body_paragraph(get_paragraph(document, paragraph_index))
    clear_paragraph(get_paragraph(document, 43))
    set_paragraph_text_like_template(get_paragraph(document, 44), str(front["en_keywords"]))
    format_body_paragraph(get_paragraph(document, 44))


def format_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.0)
    paragraph.paragraph_format.line_spacing = 1.5


def format_caption_paragraph(paragraph: Paragraph) -> None:
    if "表格题注" in [style.name for style in paragraph.part.document.styles]:
        paragraph.style = "表格题注"
    else:
        paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.2


def strip_heading_style_numbering(document: Document) -> None:
    for style_name in ["Heading 1", "二级标题", "三级标题", "四级标题"]:
        if style_name not in [style.name for style in document.styles]:
            continue
        style_element = document.styles[style_name]._element
        for num_pr in style_element.xpath(".//w:pPr/w:numPr"):
            num_pr.getparent().remove(num_pr)


def enable_update_fields(document: Document) -> None:
    settings_element = document.settings.element
    existing = settings_element.xpath("./w:updateFields")
    if existing:
        existing[0].set(qn("w:val"), "true")
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_element.append(update_fields)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    if text:
        paragraph.add_run(text)


def add_heading(anchor: Paragraph, text: str, style: str, page_break_before: bool = False) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style)
    set_paragraph_text(paragraph, text)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.page_break_before = page_break_before
    return paragraph


def add_normal_paragraph(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, "Normal")
    set_paragraph_text(paragraph, text)
    format_body_paragraph(paragraph)
    return paragraph


def add_figure(anchor: Paragraph, figure: dict[str, object], figure_dir: Path) -> Paragraph:
    figure_path = figure_dir / str(figure["file"])
    if not figure_path.exists():
        raise FileNotFoundError(f"未找到插图文件: {figure_path}")

    current_anchor = add_normal_paragraph(anchor, str(figure["intro"]))

    picture_paragraph = insert_paragraph_after(current_anchor, "Normal")
    clear_paragraph(picture_paragraph)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.first_line_indent = Cm(0)
    picture_paragraph.paragraph_format.line_spacing = 1.0
    picture_run = picture_paragraph.add_run()
    picture_run.add_picture(str(figure_path), width=Cm(float(figure["width_cm"])))

    caption_paragraph = insert_paragraph_after(picture_paragraph, "Normal")
    set_paragraph_text(caption_paragraph, str(figure["caption"]))
    format_caption_paragraph(caption_paragraph)
    return caption_paragraph


def set_cell_text(cell: _Cell, text: str) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(clean_text(text))
    for extra in cell.paragraphs[1:]:
        remove_paragraph(extra)


def set_cell_text_like_template(cell: _Cell, text: str) -> None:
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    set_paragraph_text_like_template(paragraph, text)
    for extra in cell.paragraphs[1:]:
        remove_paragraph(extra)


def add_table(anchor: Paragraph, rows: list[list[str]]) -> Paragraph:
    cols = max(len(row) for row in rows)
    section = anchor.part.document.sections[-1]
    table_width = section.page_width - section.left_margin - section.right_margin
    table = anchor._parent.add_table(rows=len(rows), cols=cols, width=table_width)
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        for c_index in range(cols):
            value = row[c_index] if c_index < len(row) else ""
            set_cell_text(table.cell(r_index, c_index), value)
    anchor._p.addnext(table._tbl)
    spacer = insert_paragraph_after_element(table._tbl, anchor._parent, "Normal")
    clear_paragraph(spacer)
    return spacer


def rebuild_body(document: Document, blocks: list[dict[str, object]], figure_dir: Path | None = None) -> None:
    toc_title = get_paragraph(document, 45)
    section_break_paragraph = get_paragraph(document, 120)
    add_toc(toc_title, section_break_paragraph)

    body_anchor = section_break_paragraph
    current_anchor = body_anchor
    first_h1 = True
    for block in blocks:
        block_type = str(block["type"])
        if block_type == "heading1":
            current_anchor = add_heading(
                current_anchor,
                str(block["text"]),
                "Heading 1",
                page_break_before=not first_h1,
            )
            first_h1 = False
        elif block_type == "heading2":
            current_anchor = add_heading(current_anchor, str(block["text"]), "二级标题")
        elif block_type == "heading3":
            current_anchor = add_heading(current_anchor, str(block["text"]), "三级标题")
        elif block_type == "paragraph":
            current_anchor = add_normal_paragraph(current_anchor, str(block["text"]))
        elif block_type == "table":
            current_anchor = add_table(current_anchor, list(block["rows"]))

        heading_text = str(block.get("text", ""))
        if figure_dir is not None and block_type in {"heading2", "heading3"} and heading_text in FIGURE_INSERTS:
            for figure in FIGURE_INSERTS[heading_text]:
                current_anchor = add_figure(current_anchor, figure, figure_dir)


def remove_sample_body(document: Document) -> None:
    section_break_paragraph = get_paragraph(document, 120)
    current = section_break_paragraph._p.getnext()
    while current is not None:
        next_node = current.getnext()
        if current.tag == qn("w:sectPr"):
            break
        current.getparent().remove(current)
        current = next_node


def normalize_sections(document: Document, title: str) -> None:
    for section in document.sections[:-1]:
        section.header.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        if section.header.paragraphs:
            clear_paragraph(section.header.paragraphs[0])
        if section.first_page_header.paragraphs:
            clear_paragraph(section.first_page_header.paragraphs[0])

    section = document.sections[-1]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.header.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    if section.header.paragraphs:
        set_paragraph_text_like_template(section.header.paragraphs[0], title)
    if section.first_page_header.paragraphs:
        set_paragraph_text_like_template(section.first_page_header.paragraphs[0], title)


def main() -> None:
    args = parse_args()
    template_path = Path(args.template)
    source_path = Path(args.source)
    output_path = Path(args.output)
    figure_dir = Path(args.figure_dir) if args.figure_dir else None

    lines = read_markdown(source_path)
    front = extract_front_matter(lines)
    blocks = parse_body_blocks(lines)

    document = Document(str(template_path))
    strip_heading_style_numbering(document)
    set_cover(front, document)
    build_abstract(front, document)
    remove_sample_body(document)
    rebuild_body(document, blocks, figure_dir)
    normalize_sections(document, str(front["title"]))
    enable_update_fields(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    main()
